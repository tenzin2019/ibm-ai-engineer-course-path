"""
Document loaders for different file types (PDF, DOCX, Excel)
"""
import os
from typing import List
from langchain.schema import Document
from pypdf import PdfReader
from docx import Document as DocxDocument
import pandas as pd


class DocumentLoader:
    """Handles loading of various document types"""
    
    @staticmethod
    def load_pdf(file_path: str) -> List[Document]:
        """Load and extract text from PDF files"""
        documents = []
        try:
            reader = PdfReader(file_path)
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    documents.append(
                        Document(
                            page_content=text,
                            metadata={
                                "source": file_path,
                                "page": page_num + 1,
                                "type": "pdf"
                            }
                        )
                    )
        except Exception as e:
            raise Exception(f"Error loading PDF: {str(e)}")
        
        return documents
    
    @staticmethod
    def load_docx(file_path: str) -> List[Document]:
        """Load and extract text from DOCX files"""
        documents = []
        try:
            doc = DocxDocument(file_path)
            text_content = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Combine paragraphs into chunks
            full_text = "\n".join(text_content)
            
            if full_text.strip():
                documents.append(
                    Document(
                        page_content=full_text,
                        metadata={
                            "source": file_path,
                            "type": "docx"
                        }
                    )
                )
        except Exception as e:
            raise Exception(f"Error loading DOCX: {str(e)}")
        
        return documents
    
    @staticmethod
    def load_excel(file_path: str) -> List[Document]:
        """Load and extract text from Excel files"""
        documents = []
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convert dataframe to text representation
                text_content = f"Sheet: {sheet_name}\n\n"
                text_content += df.to_string(index=False)
                
                if text_content.strip():
                    documents.append(
                        Document(
                            page_content=text_content,
                            metadata={
                                "source": file_path,
                                "sheet": sheet_name,
                                "type": "excel"
                            }
                        )
                    )
        except Exception as e:
            raise Exception(f"Error loading Excel: {str(e)}")
        
        return documents
    
    @classmethod
    def load_document(cls, file_path: str) -> List[Document]:
        """
        Load document based on file extension
        
        Args:
            file_path: Path to the document
            
        Returns:
            List of Document objects
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return cls.load_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return cls.load_docx(file_path)
        elif file_extension in ['.xlsx', '.xls']:
            return cls.load_excel(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")

